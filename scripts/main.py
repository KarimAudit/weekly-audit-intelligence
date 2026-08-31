# -*- coding: utf-8 -*-
"""
النشرة الأسبوعية للحوكمة والتدقيق الداخلي — Governance & Audit Weekly

الجمهور: موظفو تدقيق مبتدئون ومتوسطو الخبرة يبنون كفاءتهم أسبوعًا بعد أسبوع
(لا مدراء تنفيذيون يريدون ملخصًا فقط) — هذا يحدد لهجة الأقسام "الآنية" أدناه.

المحتوى التعليمي الثابت (الومضة، الكتاب، مسار الاحتراف، توازن، مصطلح
الأسبوع، مقترحات المهام، ورشة الصياغة) كله في content_library.py وليس هنا؛
هذا الملف فقط: بحث Tavily -> توليد LLM للقسم الآني -> دمج -> DOCX -> PDF -> إيميل.

بيئة التشغيل: TAVILY_API_KEY, LLM_PROVIDER, LLM_API_KEY, LLM_MODEL,
SENDER_EMAIL, SENDER_PASSWORD, RECIPIENT_EMAILS (اختياري: SMTP_SERVER,
SMTP_PORT, LLM_MAX_TOKENS). خطوط Amiri و Carlito يجب تثبيتها على CI.
"""

import os
import re
import sys
import json
import time
import logging
import smtplib
import subprocess
import requests
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from email.header import Header
from datetime import datetime
from typing import List, Dict, Any, Optional, Tuple

try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from content_library import (
    DOMAINS, DOMAIN_ICONS, RISK_ICONS, TRUSTED_DOMAINS_BY_CATEGORY,
    DOMAIN_TO_SOURCE_CATEGORIES, FURTHER_READING_DOMAINS, STANDARDS_WATCH_DOMAINS,
    get_week_number, pick_flashback, pick_book_of_week, pick_growth_corner,
    pick_wellbeing_tip, pick_term_of_week, pick_engagement_ideas, pick_report_writing_lesson,
)

logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
log = logging.getLogger("newsletter")

# ============================================================================
# ألوان وخطوط (شكل الوثيقة — بلا تغيير)
# ============================================================================
NAVY_DARK, NAVY_PRIMARY = "0B1F3A", "13294B"
GOLD_ACCENT, GOLD_LIGHT = "C9A227", "E8D9A0"
IVORY_BG, HAIRLINE = "FBF9F4", "D8D2C2"
TEXT_DARK, TEXT_MUTED, TEXT_WHITE = "1C2430", "5B6472", "FFFFFF"
RISK_COLORS = {"حرج": "8C1D1D", "عالٍ": "B5541A", "متوسط": "8A6D1A", "منخفض": "1E6B3E"}
RELEASE_TYPE_COLORS = {"معيار نهائي": "1E6B3E", "مسودة للتعليق": "8A6D1A", "إطار أو دليل جديد": "13294B", "تحديث جوهري": "5B4E8C"}
FONT_ARABIC, FONT_LATIN = "Amiri", "Carlito"

ISSUE_NO_FILE = os.path.join(os.getcwd(), ".issue_number")
# يُكتب في جذر الـ repo عمدًا (لا بجانب هذا الملف) ليطابق git add .issue_number في الـ workflow

def verify_fonts_installed():
    for font_name, apt_pkg in [(FONT_ARABIC, "fonts-hosny-amiri"), (FONT_LATIN, "fonts-crosextra-carlito")]:
        try:
            result = subprocess.run(["fc-list", f":family={font_name}"], capture_output=True, text=True, timeout=10)
            if not result.stdout.strip():
                log.warning(f"⚠ خط '{font_name}' غير مثبت: sudo apt-get install -y {apt_pkg}")
        except Exception as e:
            log.warning(f"فشل التحقق من خط '{font_name}': {e}")

def next_issue_number() -> int:
    n = 1
    try:
        if os.path.exists(ISSUE_NO_FILE):
            with open(ISSUE_NO_FILE, "r") as f:
                n = int(f.read().strip() or "0") + 1
    except Exception:
        n = 1
    try:
        with open(ISSUE_NO_FILE, "w") as f:
            f.write(str(n))
    except Exception:
        pass
    return n


# ============================================================================
# بحث Tavily
# ============================================================================
def _request_with_retries(method: str, url: str, max_retries: int = 3, backoff_base: float = 2.0, **kwargs):
    """إعادة محاولة مع backoff أسّي عند أخطاء شبكة مؤقتة أو 429/5xx."""
    last_exc = None
    for attempt in range(1, max_retries + 1):
        try:
            resp = requests.request(method, url, **kwargs)
            if resp.status_code == 429 or resp.status_code >= 500:
                last_exc = requests.HTTPError(f"HTTP {resp.status_code}")
                if attempt < max_retries:
                    time.sleep(backoff_base ** attempt)
                    continue
                resp.raise_for_status()
            return resp
        except (requests.ConnectionError, requests.Timeout) as e:
            last_exc = e
            if attempt < max_retries:
                time.sleep(backoff_base ** attempt)
    if last_exc:
        raise last_exc
    raise RuntimeError(f"فشل الطلب إلى {url}.")

def _flatten_trusted_domains(categories: List[str]) -> List[str]:
    seen = []
    for cat in categories:
        for d in TRUSTED_DOMAINS_BY_CATEGORY.get(cat, []):
            if d not in seen:
                seen.append(d)
    return seen

def _tavily_search_call(api_key: str, query: str, include_domains: Optional[List[str]] = None, max_results: int = 4) -> dict:
    payload = {
        "api_key": api_key, "query": query, "search_depth": "advanced",
        "topic": "general", "time_range": "month", "max_results": max_results, "include_answer": True,
    }
    if include_domains:
        payload["include_domains"] = include_domains
    response = _request_with_retries("POST", "https://api.tavily.com/search", json=payload, timeout=30)
    response.raise_for_status()
    return response.json()

def _format_search_results(data: dict, header: str) -> str:
    ctx = f"### {header}:\n"
    if data.get("answer"):
        ctx += f"ملخص عام: {data['answer']}\n"
    for res in data.get("results", []):
        ctx += f"- العنوان: {res.get('title','')}\n  الرابط: {res.get('url','')}\n  المحتوى: {res.get('content','')[:500]}\n"
    return ctx

def _dedicated_search_block(api_key: str, query: str, domains: List[str], label: str) -> Tuple[str, bool]:
    try:
        data = _tavily_search_call(api_key, query, include_domains=domains, max_results=4)
        if data.get("results"):
            return _format_search_results(data, f"مصادر مخصصة لقسم {label} (من: {', '.join(domains)})"), True
        log.warning(f"لا نتائج حديثة لمصادر {label}.")
        return "", False
    except Exception as e:
        log.error(f"فشل بحث {label}: {e}")
        return "", False

def search_with_tavily(domains: List[str]) -> str:
    api_key = os.getenv("TAVILY_API_KEY")
    if not api_key:
        log.error("FATAL: TAVILY_API_KEY غير موجود.")
        return ""

    all_context, any_success = [], False
    for domain in domains:
        include_domains = _flatten_trusted_domains(DOMAIN_TO_SOURCE_CATEGORIES.get(domain, []))
        query = f"أحدث التطورات والمستجدات في {domain}"
        try:
            data = _tavily_search_call(api_key, query, include_domains=include_domains, max_results=4)
            if not data.get("results"):
                log.warning(f"لا نتائج موثوقة لـ '{domain}' — بحث احتياطي عام...")
                data = _tavily_search_call(api_key, query, include_domains=None, max_results=4)
                header = f"نتائج البحث لـ {domain} (تنبيه: بحث احتياطي عام)"
            else:
                header = f"نتائج البحث لـ {domain} (من المصادر الموثوقة: {', '.join(include_domains)})"
            all_context.append(_format_search_results(data, header))
            any_success = any_success or bool(data.get("results"))
        except Exception as e:
            log.error(f"فشل البحث لمجال {domain}: {e}")
            all_context.append(f"### نتائج البحث لـ {domain}:\n(تعذر جلب النتائج)\n")

    fr_text, fr_ok = _dedicated_search_block(
        api_key, "أحدث المقالات والتقارير في التدقيق الداخلي والحوكمة وإدارة المخاطر",
        FURTHER_READING_DOMAINS, "further_reading")
    if fr_text:
        all_context.append(fr_text)
    any_success = any_success or fr_ok

    sw_text, sw_ok = _dedicated_search_block(
        api_key, "إصدار معيار جديد أو تحديث إطار عمل أو مسودة للتعليق exposure draft",
        STANDARDS_WATCH_DOMAINS, "standards_watch")
    if sw_text:
        all_context.append(sw_text)
    any_success = any_success or sw_ok

    if not any_success:
        log.error("FATAL: فشل البحث في كل المجالات — لا سياق حقيقي لإرساله للـ LLM.")
        return ""
    return "\n\n".join(all_context)


# ============================================================================
# توليد المحتوى الآني عبر LLM
# ============================================================================
def generate_with_llm(prompt: str) -> Optional[str]:
    provider = os.getenv("LLM_PROVIDER", "glm").lower()
    api_key = os.getenv("LLM_API_KEY")
    model = os.getenv("LLM_MODEL", "glm-4-flash")
    if not api_key:
        log.error("FATAL: LLM_API_KEY غير موجود.")
        return None

    base_urls = {"glm": "https://open.bigmodel.cn/api/paas/v4", "deepseek": "https://api.deepseek.com/v1", "openai": "https://api.openai.com/v1"}
    if provider not in base_urls:
        log.error(f"LLM_PROVIDER غير مدعوم: {provider}")
        return None
    url = f"{base_urls[provider]}/chat/completions"
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}

    # 10 مجالات + تحليل أعمق بلهجة تعليمية = JSON أكبر من افتراضي GLM القديم
    # (6000)؛ رُفع لـ 16000 بأمان (تحققنا أن GLM/DeepSeek يدعمان أكثر بكثير).
    max_tokens = int(os.getenv("LLM_MAX_TOKENS", "16000"))
    payload = {
        "model": model, "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.4, "max_tokens": max_tokens, "response_format": {"type": "json_object"},
    }

    response = None
    try:
        log.info(f"Generating with {provider}/{model} (max_tokens={max_tokens})...")
        response = _request_with_retries("POST", url, json=payload, headers=headers, timeout=90)
        response.raise_for_status()
        choice = response.json()["choices"][0]
        content_text = choice["message"]["content"]
        if choice.get("finish_reason") == "length":
            log.warning(f"الرد اتقطع (finish_reason=length) رغم max_tokens={max_tokens} — جرّب رفع LLM_MAX_TOKENS.")
        return content_text
    except Exception as e:
        log.error(f"FATAL: فشل استدعاء LLM ({type(e).__name__}): {e}")
        if response is not None:
            log.error(f"Response Body: {response.text[:2000]}")
        return None


def build_generation_prompt(search_context: str) -> str:
    return f"""أنت مدرّب تقني خبير (Senior Mentor) في التدقيق الداخلي والحوكمة، تكتب لموظفي تدقيق
مبتدئين ومتوسطي الخبرة — لا لمدراء تنفيذيين يريدون ملخصًا فقط. هدفك ليس إخبار القارئ بالخبر،
بل تعليمه كيف يفهمه ويستخدمه ليبني كفاءته المهنية أسبوعًا بعد أسبوع. لا تفترض معرفة مسبقة
بأي مصطلح متخصص يظهر في كلامك — عرّفه بجملة مبسطة أول مرة يظهر فيها.

القاعدة الذهبية: كل جملة يجب أن تُعلّم شيئًا أو تخبر بواقعة محددة من نتائج البحث، لا أن
تكون حشوًا عامًا يصح كتابته دون قراءة أي بحث.

=== نتائج بحث فعلية وحديثة — المصدر الوحيد المسموح به للوقائع ===
{search_context}
====================================================================

ممنوع اختلاق أخبار/أرقام/روابط/تواريخ غير موجودة حرفيًا أعلاه. لم تجد معلومة كافية لحقل؟
أعده فارغًا ({{}} أو []) بدل الاختلاق.

غطِّ المجالات التالية، عنصر واحد لكل مجال قدر الإمكان (استثنِ فقط ما لا سياق كافٍ له):
{chr(10).join(f'- {d}' for d in DOMAINS)}

=== معايير الجودة ===
issue_theme: جملة واحدة تلخص الخيط الناظم بين أبرز 2-3 تطورات هذا الأسبوع.

editor_note (3-5 أسطر): افتتاحية ودّية بصوت مدرّب لا مدير — تربط بين تطورين لتكشف نمطًا،
وتشرح للقارئ الأقل خبرة لماذا يستحق انتباهه. أنهِ بفكرة تستحق تأمله.

exec_brief (3-5 نقاط): "لمحة الأسبوع" — كل نقطة جملة واحدة تلخص أهم تطور برقم أو واقعة محددة.

stat_of_week: رقم واحد حقيقي 100% من البحث مع سياق يفسّر أهميته في جملتين. لا رقم حقيقي
واضح؟ كائن فارغ {{}}.

standards_watch (0-3 عناصر): معيار/مسودة/دليل جديد فعليًا من الجهات الموثوقة. صنّف
release_type (معيار نهائي | مسودة للتعليق | إطار أو دليل جديد | تحديث جوهري)، واذكر
comment_deadline أو effective_date إن وُجدا. download_url رابط حقيقي من البحث يقود
للوثيقة أو صفحة الوصول — هذا الحقل هو سبب وجود القسم، لا تتركه فارغًا إن وُجد رابط حقيقي.

domain_updates (قلب النشرة، موجَّه لقارئ يتعلّم لا يراجع سريعًا): لكل عنصر —
  - summary (4-5 جمل): ماذا حدث فعليًا بتفاصيل محددة من البحث، مع شرح مبسّط لأي مصطلح أو
    إطار مذكور قد لا يعرفه قارئ متوسط الخبرة (بجملة قصيرة بين قوسين مثلاً).
  - why_it_matters: لماذا يهم هذا تحديدًا لمن يبني خبرته المهنية الآن، لا لعضو مجلس إدارة.
  - recommended_actions: ما الذي يمكن لموظف تدقيق أن يفعله عمليًا هذا الأسبوع بهذه المعلومة
    (يقرأ الوثيقة الكاملة؟ يراجع ملفًا مشابهًا في عمله؟ يسأل مديره عن كذا؟).
  - skill_gained: المهارة أو المعرفة المحددة التي يكتسبها القارئ من فهم هذا التحديث
    (3-6 كلمات، مثل "تحليل تقارير COSO" أو "فهم Exposure Drafts").

further_reading: فقط من قسم "مصادر مخصصة لقسم further_reading" في البحث. غير موجود؟ [].

أعد JSON صِرف فقط (بدون نص قبله/بعده، بدون fences) بهذا الهيكل تمامًا:

{{
  "issue_theme": "...", "editor_note": "...", "exec_brief": ["...", "...", "..."],
  "stat_of_week": {{"value": "...", "context_ar": "...", "source_name": "...", "source_url": "..."}},
  "standards_watch": [{{"title": "...", "issuing_body": "...", "release_type": "...", "issue_date": "...", "effective_date": "...", "comment_deadline": "...", "summary_ar": "...", "impact_ar": "...", "download_url": "...", "source_name": "..."}}],
  "domain_updates": [{{"domain": "...", "title": "...", "category": "...", "risk_level": "حرج | عالٍ | متوسط | منخفض", "skill_gained": "...", "summary": "...", "why_it_matters": "...", "recommended_actions": "...", "source_name": "...", "source_url": "..."}}],
  "further_reading": [{{"title": "...", "source_name": "...", "url": "...", "type": "مقال | تقرير | دراسة"}}]
}}

عربية فصحى واضحة (لا متعالية)، كل مصطلح تقني بالإنجليزية بين قوسين أول ورود له. اشرح، لا
تلخّص فقط — القارئ يتعلم منك، لا يراجع معك.
"""


def _extract_json(text: str) -> Optional[dict]:
    if not text:
        return None
    cleaned = re.sub(r"^```(?:json)?\s*|\s*```$", "", text.strip(), flags=re.MULTILINE)
    match = re.search(r"\{.*\}", cleaned, flags=re.DOTALL)
    try:
        return json.loads(match.group(0) if match else cleaned)
    except json.JSONDecodeError as e:
        log.error(f"JSON parse failure: {e}")
        return None

def _validate_content_shape(data: dict) -> bool:
    required = ["issue_theme", "editor_note", "exec_brief", "domain_updates", "further_reading"]
    missing = [k for k in required if k not in data]
    if missing:
        log.error(f"استجابة LLM ناقصة: {missing}")
        return False
    if not isinstance(data.get("domain_updates"), list) or not data["domain_updates"]:
        log.error("domain_updates فاضي — لا فائدة من نشرة بدون تحديثات.")
        return False
    return True


def generate_newsletter_content() -> Tuple[Optional[Dict[str, Any]], bool]:
    search_context = search_with_tavily(DOMAINS)
    if not search_context:
        log.error("Pipeline STOPPED: تعذر جلب نتائج البحث.")
        return None, False

    raw_response = generate_with_llm(build_generation_prompt(search_context))
    if not raw_response:
        return None, False

    data = _extract_json(raw_response)
    if not data:
        log.error(f"FATAL: تعذر تحليل JSON. أول 2000 حرف: {raw_response[:2000]}")
        return None, False

    data.setdefault("stat_of_week", {})
    data.setdefault("standards_watch", [])
    data.setdefault("exec_brief", [])
    data.setdefault("further_reading", [])

    # الأركان التعليمية الثابتة — من content_library.py، بلا أي تدخل من الـ LLM
    week = get_week_number()
    data["flashback"] = pick_flashback(week)
    data["book_of_week"] = pick_book_of_week(week)
    data["growth_corner"] = pick_growth_corner(week)
    data["wellbeing_corner"] = pick_wellbeing_tip(week)
    data["term_of_week"] = pick_term_of_week(week)
    data["engagement_ideas"] = pick_engagement_ideas(week)
    data["report_writing_lesson"] = pick_report_writing_lesson(week)

    if not _validate_content_shape(data):
        log.error(f"Raw response last 500 chars: ...{raw_response[-500:]}")
        return None, False

    log.info(f"OK: {len(data['domain_updates'])} domain updates, {len(data['standards_watch'])} standards items.")
    return data, True


# ============================================================================
# أدوات DOCX
# ============================================================================
def set_rtl(paragraph):
    paragraph._p.get_or_add_pPr().append(OxmlElement("w:bidi"))

def style_run(run, size=11, bold=False, color=TEXT_DARK, italic=False):
    run.font.size, run.font.bold, run.font.italic = Pt(size), bold, italic
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = FONT_LATIN
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi"):
        rFonts.set(qn(attr), FONT_LATIN)
    for attr in ("w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), FONT_ARABIC)
    return run

def add_hyperlink(paragraph, text, url, color=GOLD_ACCENT, underline=True, size=10.5, bold=False):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink"); hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), FONT_LATIN); rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    rFonts.set(qn("w:cs"), FONT_ARABIC); rFonts.set(qn("w:eastAsia"), FONT_ARABIC)
    rPr.append(rFonts)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(size * 2))); rPr.append(sz)
    if bold:
        rPr.append(OxmlElement("w:b"))
    if underline:
        u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    color_el = OxmlElement("w:color"); color_el.set(qn("w:val"), color); rPr.append(color_el)
    new_run.append(rPr)
    t = OxmlElement("w:t"); t.text = text; new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def set_cell_background(cell, color_hex: str):
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'))

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcMar = OxmlElement("w:tcMar")
    for name, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        node = OxmlElement(f"w:{name}"); node.set(qn("w:w"), str(val)); node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    cell._tc.get_or_add_tcPr().append(tcMar)

def add_gold_rule(doc, space_before=2, space_after=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(space_before), Pt(space_after)
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), GOLD_ACCENT)
    pBdr.append(bottom); p._p.get_or_add_pPr().append(pBdr)
    return p

def _boxed_cell(doc, bg=IVORY_BG, top=120, bottom=120, left=200, right=200, border_color=None, border_side="right"):
    """جدول بخلية واحدة كصندوق تصميم — نمط مكرر عبر أغلب الأقسام."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, bg)
    set_cell_margins(cell, top=top, bottom=bottom, left=left, right=right)
    if border_color:
        tcPr = cell._tc.get_or_add_tcPr()
        tcPr.append(parse_xml(f'<w:tcBorders {nsdecls("w")}><w:{border_side} w:val="single" w:sz="20" w:space="0" w:color="{border_color}"/></w:tcBorders>'))
    return cell


# ============================================================================
# بناء DOCX — الأقسام
# ============================================================================
def add_cover_page(doc, content: Dict[str, Any], issue_no: int):
    cell = _boxed_cell(doc, bg=NAVY_DARK, top=900, bottom=900, left=500, right=500)
    p_kicker = cell.paragraphs[0]; p_kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER; set_rtl(p_kicker)
    style_run(p_kicker.add_run(f"العدد {issue_no}  •  {datetime.now().strftime('%Y-%m-%d')}"), size=11, color=GOLD_LIGHT, bold=True)

    p_title = cell.add_paragraph(); p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER; set_rtl(p_title)
    p_title.paragraph_format.space_before = Pt(14)
    style_run(p_title.add_run("النشرة الأسبوعية للحوكمة والتدقيق الداخلي"), size=26, bold=True, color=TEXT_WHITE)

    p_theme = cell.add_paragraph(); p_theme.alignment = WD_ALIGN_PARAGRAPH.CENTER; set_rtl(p_theme)
    p_theme.paragraph_format.space_before = Pt(10)
    style_run(p_theme.add_run(content.get("issue_theme", "")), size=14, italic=True, color=GOLD_ACCENT)

    p_sub = cell.add_paragraph(); p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER; set_rtl(p_sub)
    p_sub.paragraph_format.space_before = Pt(16)
    style_run(p_sub.add_run(" • ".join(DOMAINS)), size=8.5, color=GOLD_LIGHT)
    doc.add_paragraph()

def add_editor_note(doc, content: Dict[str, Any]):
    p = doc.add_paragraph(); set_rtl(p); p.paragraph_format.space_after = Pt(4)
    style_run(p.add_run("✉️ كلمة التحرير"), size=13, bold=True, color=NAVY_PRIMARY)
    p2 = doc.add_paragraph(); set_rtl(p2); p2.paragraph_format.line_spacing = 1.3
    style_run(p2.add_run(content.get("editor_note", "")), size=11.5, color=TEXT_MUTED, italic=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_exec_brief(doc, brief: List[str]):
    if not brief:
        return
    header = doc.add_paragraph(); set_rtl(header)
    style_run(header.add_run("⚡ لمحة الأسبوع"), size=13, bold=True, color=NAVY_PRIMARY)
    cell = _boxed_cell(doc, bg=NAVY_PRIMARY, top=140, bottom=140, left=220, right=220)
    for i, point in enumerate(brief):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        set_rtl(p); p.paragraph_format.space_after = Pt(6)
        style_run(p.add_run("◆ "), size=11, bold=True, color=GOLD_ACCENT)
        style_run(p.add_run(point), size=11, color=TEXT_WHITE)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

def add_executive_dashboard(doc, updates: List[Dict[str, Any]], standards_count: int):
    tbl = doc.add_table(rows=2, cols=4); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    total = len(updates)
    critical_high = sum(1 for b in updates if b.get("risk_level") in ("حرج", "عالٍ"))
    domains_covered = len(set(b.get("domain", "") for b in updates))
    headers = ["إجمالي التحديثات", "معايير جديدة", "أولوية عالية", "المجالات المغطاة"]
    values = [str(total), str(standards_count), str(critical_high), str(domains_covered)]
    for i, (h, v) in enumerate(zip(headers, values)):
        top_cell = tbl.cell(0, i); set_cell_background(top_cell, NAVY_PRIMARY); set_cell_margins(top_cell, top=100, bottom=80)
        p = top_cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; set_rtl(p)
        style_run(p.add_run(h), size=9.5, bold=True, color=TEXT_WHITE)
        bottom_cell = tbl.cell(1, i); set_cell_background(bottom_cell, IVORY_BG); set_cell_margins(bottom_cell, top=140, bottom=140)
        p2 = bottom_cell.paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_run(p2.add_run(v), size=18, bold=True, color=(GOLD_ACCENT if i != 2 else RISK_COLORS["حرج"]))
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

def add_stat_of_week(doc, stat: Dict[str, Any]):
    if not stat or not stat.get("value"):
        return
    header = doc.add_paragraph(); set_rtl(header)
    style_run(header.add_run("🔢 رقم الأسبوع"), size=13, bold=True, color=NAVY_PRIMARY)
    cell = _boxed_cell(doc, bg=IVORY_BG, top=160, bottom=160, left=220, right=220, border_color=GOLD_ACCENT)
    p_val = cell.paragraphs[0]; set_rtl(p_val)
    style_run(p_val.add_run(stat.get("value", "")), size=28, bold=True, color=NAVY_PRIMARY)
    p_ctx = cell.add_paragraph(); set_rtl(p_ctx); p_ctx.paragraph_format.space_before = Pt(6); p_ctx.paragraph_format.line_spacing = 1.3
    style_run(p_ctx.add_run(stat.get("context_ar", "")), size=11, color=TEXT_DARK)
    if stat.get("source_url"):
        p_src = cell.add_paragraph(); set_rtl(p_src); p_src.paragraph_format.space_before = Pt(6)
        style_run(p_src.add_run("المصدر: "), size=9.5, bold=True, color=TEXT_MUTED)
        add_hyperlink(p_src, stat.get("source_name", stat["source_url"]), stat["source_url"], size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

def add_standards_watch(doc, items: List[Dict[str, Any]]):
    if not items:
        return
    header = doc.add_paragraph(); set_rtl(header)
    style_run(header.add_run("🆕 رادار المعايير والإصدارات الجديدة"), size=13, bold=True, color=NAVY_PRIMARY)
    add_gold_rule(doc, space_before=2, space_after=8)
    for item in items:
        release_type = item.get("release_type", "تحديث جوهري")
        type_color = RELEASE_TYPE_COLORS.get(release_type, NAVY_PRIMARY)
        cell = _boxed_cell(doc, bg="FFFFFF", top=140, bottom=140, border_color=type_color)
        tcPr = cell._tc.get_or_add_tcPr()
        tcPr.append(parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="single" w:sz="4" w:space="0" w:color="{HAIRLINE}"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="{HAIRLINE}"/><w:left w:val="single" w:sz="4" w:space="0" w:color="{HAIRLINE}"/></w:tcBorders>'))

        p_head = cell.paragraphs[0]; set_rtl(p_head)
        style_run(p_head.add_run(item.get("title", "")), size=12.5, bold=True, color=NAVY_PRIMARY)
        style_run(p_head.add_run(f"  [{item.get('issuing_body','')} | {release_type}]"), size=9.5, bold=True, color=type_color)

        dates = []
        if item.get("issue_date"):
            dates.append(f"الإصدار: {item['issue_date']}")
        if item.get("effective_date"):
            dates.append(f"السريان: {item['effective_date']}")
        if item.get("comment_deadline"):
            dates.append(f"آخر أجل للتعليق: {item['comment_deadline']}")
        if dates:
            p_d = cell.add_paragraph(); set_rtl(p_d); p_d.paragraph_format.space_before = Pt(3)
            style_run(p_d.add_run("  |  ".join(dates)), size=9, italic=True, color=TEXT_MUTED)

        p_sum = cell.add_paragraph(); set_rtl(p_sum); p_sum.paragraph_format.space_before = Pt(6); p_sum.paragraph_format.line_spacing = 1.3
        style_run(p_sum.add_run(item.get("summary_ar", "")), size=10.5, color=TEXT_DARK)

        if item.get("impact_ar"):
            p_imp = cell.add_paragraph(); set_rtl(p_imp); p_imp.paragraph_format.space_before = Pt(4)
            style_run(p_imp.add_run("لماذا يهمكم: "), size=10, bold=True, color=NAVY_PRIMARY)
            style_run(p_imp.add_run(item["impact_ar"]), size=10, color=TEXT_DARK)

        if item.get("download_url"):
            p_dl = cell.add_paragraph(); set_rtl(p_dl); p_dl.paragraph_format.space_before = Pt(8)
            style_run(p_dl.add_run("⬇ "), size=11, bold=True, color=GOLD_ACCENT)
            add_hyperlink(p_dl, f"تحميل / الاطلاع على {item.get('source_name', 'الوثيقة الكاملة')}", item["download_url"], size=11, bold=True)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

def add_engagement_ideas(doc, ideas: List[Dict[str, Any]]):
    """مقترحات مهام تدقيق (core business / استراتيجية / KPIs) — من content_library."""
    if not ideas:
        return
    header = doc.add_paragraph(); set_rtl(header)
    style_run(header.add_run("🧭 مقترحات لمهام تدقيق"), size=13, bold=True, color=NAVY_PRIMARY)
    sub = doc.add_paragraph(); set_rtl(sub); sub.paragraph_format.space_after = Pt(6)
    style_run(sub.add_run("أفكار جاهزة لخطة عملكم القادمة — عناصر أساسية، لا مجرد عناوين"), size=9.5, italic=True, color=TEXT_MUTED)
    for idea in ideas:
        cell = _boxed_cell(doc, bg="FFFFFF", top=130, bottom=130, border_color=NAVY_PRIMARY)
        p_t = cell.paragraphs[0]; set_rtl(p_t)
        style_run(p_t.add_run(idea.get("title", "")), size=12, bold=True, color=NAVY_PRIMARY)

        p_o = cell.add_paragraph(); set_rtl(p_o); p_o.paragraph_format.space_before = Pt(4)
        style_run(p_o.add_run("الهدف: "), size=10, bold=True, color=NAVY_PRIMARY)
        style_run(p_o.add_run(idea.get("objective_ar", "")), size=10, color=TEXT_DARK)

        if idea.get("scope_questions"):
            p_q = cell.add_paragraph(); set_rtl(p_q); p_q.paragraph_format.space_before = Pt(4)
            style_run(p_q.add_run("أسئلة النطاق الرئيسية:"), size=10, bold=True, color=NAVY_PRIMARY)
            for q in idea["scope_questions"]:
                pq = cell.add_paragraph(); set_rtl(pq); pq.paragraph_format.space_before = Pt(2)
                style_run(pq.add_run("• "), size=10, bold=True, color=GOLD_ACCENT)
                style_run(pq.add_run(q), size=10, color=TEXT_DARK)

        if idea.get("watch_for_ar"):
            p_w = cell.add_paragraph(); set_rtl(p_w); p_w.paragraph_format.space_before = Pt(5)
            style_run(p_w.add_run("🚩 مؤشرات إنذار: "), size=9.5, bold=True, color=RISK_COLORS["عالٍ"])
            style_run(p_w.add_run(idea["watch_for_ar"]), size=9.5, italic=True, color=TEXT_MUTED)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

def add_report_writing_lesson(doc, lesson: Dict[str, Any]):
    """ورشة صياغة تقرير التدقيق — مبدأ + مثال ضعيف + مثال قوي، من content_library."""
    if not lesson or not lesson.get("principle_ar"):
        return
    header = doc.add_paragraph(); set_rtl(header)
    style_run(header.add_run("✍️ ورشة الصياغة"), size=13, bold=True, color=NAVY_PRIMARY)
    sub = doc.add_paragraph(); set_rtl(sub); sub.paragraph_format.space_after = Pt(4)
    style_run(sub.add_run(f"هذا الأسبوع: {lesson.get('category','')}"), size=10, italic=True, color=TEXT_MUTED)

    cell = _boxed_cell(doc, bg=IVORY_BG, top=140, bottom=140, border_color=GOLD_ACCENT)
    p1 = cell.paragraphs[0]; set_rtl(p1); p1.paragraph_format.line_spacing = 1.3
    style_run(p1.add_run(lesson.get("principle_ar", "")), size=10.5, color=TEXT_DARK)

    p2 = cell.add_paragraph(); set_rtl(p2); p2.paragraph_format.space_before = Pt(8)
    style_run(p2.add_run("❌ صياغة ضعيفة: "), size=10, bold=True, color=RISK_COLORS["حرج"])
    style_run(p2.add_run(lesson.get("weak_example_ar", "")), size=10, italic=True, color=TEXT_MUTED)

    p3 = cell.add_paragraph(); set_rtl(p3); p3.paragraph_format.space_before = Pt(6)
    style_run(p3.add_run("✅ صياغة قوية: "), size=10, bold=True, color=RISK_COLORS["منخفض"])
    style_run(p3.add_run(lesson.get("strong_example_ar", "")), size=10, color=TEXT_DARK)

    p4 = cell.add_paragraph(); set_rtl(p4); p4.paragraph_format.space_before = Pt(6)
    style_run(p4.add_run("لماذا الفرق مهم: "), size=9.5, bold=True, color=NAVY_PRIMARY)
    style_run(p4.add_run(lesson.get("why_ar", "")), size=9.5, color=TEXT_DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

def add_flashback(doc, flashback: Dict[str, Any]):
    if not flashback or not flashback.get("content_ar"):
        return
    header = doc.add_paragraph(); set_rtl(header)
    style_run(header.add_run(f"🧠 ومضة إدارية — {flashback.get('topic', '')}"), size=13, bold=True, color=NAVY_PRIMARY)
    cell = _boxed_cell(doc, bg="FFFFFF", top=120, bottom=120)
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="single" w:sz="4" w:space="0" w:color="{HAIRLINE}"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="{HAIRLINE}"/><w:left w:val="single" w:sz="4" w:space="0" w:color="{HAIRLINE}"/><w:right w:val="single" w:sz="4" w:space="0" w:color="{HAIRLINE}"/></w:tcBorders>'))
    paragraphs = flashback.get("content_ar", "").split("\n\n")
    for i, para in enumerate(paragraphs):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        set_rtl(p); p.paragraph_format.line_spacing = 1.3
        if i > 0:
            p.paragraph_format.space_before = Pt(6)
        style_run(p.add_run(para), size=11, color=TEXT_DARK)
    for kp in flashback.get("key_points", []):
        pk = cell.add_paragraph(); set_rtl(pk); pk.paragraph_format.space_before = Pt(4)
        style_run(pk.add_run("— "), size=10.5, bold=True, color=GOLD_ACCENT)
        style_run(pk.add_run(kp), size=10.5, color=TEXT_MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

def add_term_of_week(doc, term: Dict[str, str]):
    if not term or not term.get("term_en"):
        return
    header = doc.add_paragraph(); set_rtl(header)
    style_run(header.add_run("📖 مصطلح الأسبوع"), size=13, bold=True, color=NAVY_PRIMARY)
    cell = _boxed_cell(doc, bg=NAVY_PRIMARY, top=120, bottom=120)
    p_term = cell.paragraphs[0]; set_rtl(p_term)
    style_run(p_term.add_run(f"{term.get('term_ar', '')}  "), size=13, bold=True, color=TEXT_WHITE)
    style_run(p_term.add_run(f"({term.get('term_en', '')})"), size=11, italic=True, color=GOLD_LIGHT)
    p_def = cell.add_paragraph(); set_rtl(p_def); p_def.paragraph_format.space_before = Pt(6); p_def.paragraph_format.line_spacing = 1.3
    style_run(p_def.add_run(term.get("definition_ar", "")), size=10.5, color=TEXT_WHITE)
    if term.get("example_ar"):
        p_ex = cell.add_paragraph(); set_rtl(p_ex); p_ex.paragraph_format.space_before = Pt(6)
        style_run(p_ex.add_run("في الممارسة: "), size=9.5, bold=True, color=GOLD_LIGHT)
        style_run(p_ex.add_run(term["example_ar"]), size=9.5, italic=True, color=GOLD_LIGHT)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

def create_section_header(doc, title: str, risk: str, category: str, skill_gained: str = "", icon: str = ""):
    cell = _boxed_cell(doc, bg=NAVY_PRIMARY, top=130, bottom=130)
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; set_rtl(p)
    label = f"{icon} {title}".strip()
    style_run(p.add_run(f"{label}  "), size=13.5, bold=True, color=TEXT_WHITE)
    style_run(p.add_run(f" {RISK_ICONS.get(risk, '')} [{category} | {risk}]"), size=10, bold=True, color=GOLD_LIGHT)
    if skill_gained:
        p2 = cell.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT; set_rtl(p2); p2.paragraph_format.space_before = Pt(2)
        style_run(p2.add_run(f"🎓 المهارة المكتسبة: {skill_gained}"), size=8.5, italic=True, color=GOLD_LIGHT)

def add_domain_card(doc, block: Dict[str, Any]):
    domain = block.get("domain", block.get("title", "تحديث"))
    icon = DOMAIN_ICONS.get(domain, "🆕")
    create_section_header(doc, domain, risk=block.get("risk_level", "متوسط"), category=block.get("category", "عام"),
                           skill_gained=block.get("skill_gained", ""), icon=icon)
    if block.get("title"):
        p_title = doc.add_paragraph(); set_rtl(p_title); p_title.paragraph_format.space_before = Pt(4)
        style_run(p_title.add_run(block["title"]), size=12, bold=True, color=NAVY_PRIMARY)

    p_body = doc.add_paragraph(); set_rtl(p_body); p_body.paragraph_format.line_spacing = 1.3; p_body.paragraph_format.space_after = Pt(6)
    style_run(p_body.add_run(block.get("summary", "")), size=11, color=TEXT_DARK)

    cell = _boxed_cell(doc, bg=IVORY_BG, top=110, bottom=110, left=180, right=180,
                        border_color=RISK_COLORS.get(block.get("risk_level"), GOLD_ACCENT))
    p1 = cell.paragraphs[0]; set_rtl(p1)
    style_run(p1.add_run("الأثر: "), size=10.5, bold=True, color=NAVY_PRIMARY)
    style_run(p1.add_run(block.get("why_it_matters", "")), size=10.5, color=TEXT_DARK)
    p2 = cell.add_paragraph(); set_rtl(p2); p2.paragraph_format.space_before = Pt(4)
    style_run(p2.add_run("ماذا تفعل بهذا الآن: "), size=10.5, bold=True, color=NAVY_PRIMARY)
    style_run(p2.add_run(block.get("recommended_actions", "")), size=10.5, color=TEXT_DARK)
    if block.get("source_url"):
        p3 = cell.add_paragraph(); set_rtl(p3); p3.paragraph_format.space_before = Pt(4)
        style_run(p3.add_run("المصدر: "), size=9.5, bold=True, color=TEXT_MUTED)
        add_hyperlink(p3, block.get("source_name", block["source_url"]), block["source_url"], size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(10)

def add_growth_wellbeing(doc, growth: Dict[str, str], wellbeing: Dict[str, str]):
    if growth and growth.get("insight_ar"):
        header = doc.add_paragraph(); set_rtl(header)
        style_run(header.add_run("🌱 مسار الاحتراف — كيف يفكّر الكبار"), size=13, bold=True, color=NAVY_PRIMARY)
        cell = _boxed_cell(doc, bg="FFFFFF", top=120, bottom=120, border_color=NAVY_PRIMARY)
        tcPr = cell._tc.get_or_add_tcPr()
        tcPr.append(parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="single" w:sz="4" w:space="0" w:color="{HAIRLINE}"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="{HAIRLINE}"/><w:left w:val="single" w:sz="4" w:space="0" w:color="{HAIRLINE}"/></w:tcBorders>'))
        p1 = cell.paragraphs[0]; set_rtl(p1); p1.paragraph_format.line_spacing = 1.3
        style_run(p1.add_run(growth.get("insight_ar", "")), size=11, color=TEXT_DARK)
        p2 = cell.add_paragraph(); set_rtl(p2); p2.paragraph_format.space_before = Pt(5)
        style_run(p2.add_run("كيف تطبّقها هذا الأسبوع: "), size=10, bold=True, color=NAVY_PRIMARY)
        style_run(p2.add_run(growth.get("how_to_apply_ar", "")), size=10, italic=True, color=TEXT_MUTED)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    if wellbeing and wellbeing.get("tip_ar"):
        header2 = doc.add_paragraph(); set_rtl(header2)
        style_run(header2.add_run("⚖️ لحظة توازن"), size=13, bold=True, color=NAVY_PRIMARY)
        cell2 = _boxed_cell(doc, bg=IVORY_BG, top=110, bottom=110)
        p = cell2.paragraphs[0]; set_rtl(p); p.paragraph_format.line_spacing = 1.3
        style_run(p.add_run(wellbeing.get("tip_ar", "")), size=10.5, color=TEXT_DARK, italic=True)
        doc.add_paragraph().paragraph_format.space_after = Pt(8)

def add_further_reading(doc, items: List[Dict[str, str]], book: Dict[str, str]):
    if not items and not (book and book.get("title")):
        return
    header = doc.add_paragraph(); set_rtl(header)
    style_run(header.add_run("📚 للقراءة هذا الأسبوع"), size=13, bold=True, color=NAVY_PRIMARY)
    for it in items:
        p = doc.add_paragraph(); set_rtl(p); p.paragraph_format.space_after = Pt(4)
        style_run(p.add_run(f"• {it.get('type', 'مقال')} — "), size=10.5, color=TEXT_MUTED)
        if it.get("url"):
            add_hyperlink(p, it.get("title", ""), it["url"], size=10.5, bold=True)
        else:
            style_run(p.add_run(it.get("title", "")), size=10.5, bold=True, color=TEXT_DARK)
        if it.get("source_name"):
            style_run(p.add_run(f"  ({it['source_name']})"), size=9.5, color=TEXT_MUTED)
    if book and book.get("title"):
        add_gold_rule(doc, space_before=8, space_after=6)
        p = doc.add_paragraph(); set_rtl(p)
        style_run(p.add_run("📕 كتاب الأسبوع: "), size=11, bold=True, color=NAVY_PRIMARY)
        style_run(p.add_run(f"{book['title']} — {book.get('author', '')}"), size=11, bold=True, color=TEXT_DARK)
        p2 = doc.add_paragraph(); set_rtl(p2)
        style_run(p2.add_run(book.get("why_ar", "")), size=10.5, italic=True, color=TEXT_MUTED)
        if book.get("key_takeaway_ar"):
            p3 = doc.add_paragraph(); set_rtl(p3); p3.paragraph_format.space_before = Pt(4)
            style_run(p3.add_run("💡 طبّقه هذا الأسبوع: "), size=10, bold=True, color=NAVY_PRIMARY)
            style_run(p3.add_run(book["key_takeaway_ar"]), size=10, color=TEXT_DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

def build_word_document(content: Dict[str, Any], reports_dir: str = "reports") -> str:
    os.makedirs(reports_dir, exist_ok=True)
    issue_no = next_issue_number()
    filename = os.path.join(reports_dir, f"Governance_Weekly_Brief_{datetime.now().strftime('%Y-%m-%d')}_Issue{issue_no}.docx")
    doc = Document()

    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(0.7)
        section.left_margin = section.right_margin = Inches(0.8)
        f_p = section.footer.paragraphs[0]; f_p.alignment = WD_ALIGN_PARAGRAPH.CENTER; set_rtl(f_p)
        style_run(f_p.add_run("النشرة الأسبوعية للحوكمة والتدقيق — سري وللاستخدام الداخلي فقط  |  "), size=8, color=TEXT_MUTED)
        fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
        fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
        run = f_p.add_run(); run._r.append(fld1); run._r.append(instr); run._r.append(fld2)
        style_run(run, size=8, color=TEXT_MUTED)

    add_cover_page(doc, content, issue_no)
    add_editor_note(doc, content)
    add_exec_brief(doc, content.get("exec_brief", []))
    add_executive_dashboard(doc, content.get("domain_updates", []), len(content.get("standards_watch", [])))
    add_stat_of_week(doc, content.get("stat_of_week", {}))
    add_standards_watch(doc, content.get("standards_watch", []))
    add_report_writing_lesson(doc, content.get("report_writing_lesson", {}))
    add_flashback(doc, content.get("flashback", {}))
    add_term_of_week(doc, content.get("term_of_week", {}))

    updates_header = doc.add_paragraph(); set_rtl(updates_header)
    style_run(updates_header.add_run("📌 أبرز التطورات حسب المجال"), size=13, bold=True, color=NAVY_PRIMARY)
    add_gold_rule(doc, space_before=2, space_after=8)
    for block in content.get("domain_updates", []):
        add_domain_card(doc, block)

    add_engagement_ideas(doc, content.get("engagement_ideas", []))
    add_growth_wellbeing(doc, content.get("growth_corner", {}), content.get("wellbeing_corner", {}))
    add_further_reading(doc, content.get("further_reading", []), content.get("book_of_week", {}))
    doc.save(filename)
    log.info(f"Generated Word document: {filename}")
    return filename


# ============================================================================
# تحويل PDF (بلا تغيير)
# ============================================================================
def convert_docx_to_pdf(docx_path: str) -> str:
    pdf_path = docx_path.rsplit(".", 1)[0] + ".pdf"
    try:
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        return pdf_path
    except Exception as e1:
        log.warning(f"docx2pdf unavailable: {e1}. Trying LibreOffice...")
    try:
        subprocess.run(["libreoffice", "--headless", "--convert-to", "pdf", docx_path, "--outdir", os.path.dirname(docx_path)],
                        check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        return pdf_path
    except Exception as e2:
        log.warning(f"LibreOffice failed: {e2}. Falling back to DOCX attachment.")
    return docx_path


# ============================================================================
# إيميل HTML
# ============================================================================
def _esc(s: str) -> str:
    return (s or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

def generate_email_html(content: Dict[str, Any], issue_no: int) -> str:
    exec_brief_html = ""
    if content.get("exec_brief"):
        pts = "".join(f'<li style="margin-bottom:8px;">{_esc(p)}</li>' for p in content["exec_brief"])
        exec_brief_html = (f'<div style="background:#13294B;border-radius:8px;padding:16px 20px;margin-bottom:22px;text-align:right;">'
                            f'<div style="color:#E8D9A0;font-weight:bold;font-size:13px;margin-bottom:8px;">⚡ لمحة الأسبوع</div>'
                            f'<ul style="margin:0;padding-right:18px;color:#fff;font-size:13px;">{pts}</ul></div>')

    stat = content.get("stat_of_week", {})
    stat_html = ""
    if stat and stat.get("value"):
        src = ""
        if stat.get("source_url"):
            src = f'<a href="{stat["source_url"]}" style="color:#5B6472;font-size:11px;">المصدر: {_esc(stat.get("source_name",""))}</a>'
        stat_html = (f'<div style="background:#FBF9F4;border-right:4px solid #C9A227;border-radius:8px;padding:16px 20px;margin-bottom:22px;text-align:right;">'
                     f'<div style="color:#13294B;font-weight:bold;font-size:13px;margin-bottom:6px;">🔢 رقم الأسبوع</div>'
                     f'<div style="color:#13294B;font-weight:bold;font-size:26px;margin-bottom:6px;">{_esc(stat.get("value",""))}</div>'
                     f'<p style="color:#1C2430;font-size:13px;line-height:1.6;margin:0 0 6px 0;">{_esc(stat.get("context_ar",""))}</p>{src}</div>')

    standards_html = ""
    for it in content.get("standards_watch", []):
        rt = it.get("release_type", "")
        tc = RELEASE_TYPE_COLORS.get(rt, "13294B")
        dates = " | ".join(f"{lbl}: {_esc(it[k])}" for k, lbl in
                           [("issue_date", "الإصدار"), ("effective_date", "السريان"), ("comment_deadline", "آخر أجل للتعليق")] if it.get(k))
        impact_html = ""
        if it.get("impact_ar"):
            impact_html = f'<p style="color:#334155;font-size:11.5px;margin:4px 0;"><strong>لماذا يهمكم:</strong> {_esc(it["impact_ar"])}</p>'
        dl_html = ""
        if it.get("download_url"):
            dl_html = (f'<a href="{it["download_url"]}" style="display:inline-block;margin-top:8px;color:#fff;background:#C9A227;'
                       f'padding:6px 14px;border-radius:4px;font-size:11.5px;font-weight:bold;text-decoration:none;">'
                       f'⬇ تحميل / الاطلاع على {_esc(it.get("source_name","الوثيقة"))}</a>')
        standards_html += (f'<div style="background:#fff;border:1px solid #D8D2C2;border-right:4px solid #{tc};border-radius:8px;padding:14px 18px;margin-bottom:14px;text-align:right;">'
                            f'<div style="font-weight:bold;color:#13294B;font-size:13px;">{_esc(it.get("title",""))} '
                            f'<span style="background:#{tc};color:#fff;padding:1px 8px;border-radius:8px;font-size:10px;font-weight:bold;margin-right:6px;">{_esc(rt)}</span></div>'
                            f'<div style="color:#5B6472;font-size:10.5px;margin:4px 0;">{_esc(it.get("issuing_body",""))} — {dates}</div>'
                            f'<p style="color:#1C2430;font-size:12.5px;line-height:1.6;margin:6px 0;">{_esc(it.get("summary_ar",""))}</p>{impact_html}{dl_html}</div>')

    lesson = content.get("report_writing_lesson", {})
    lesson_html = ""
    if lesson and lesson.get("principle_ar"):
        lesson_html = (f'<div style="background:#FBF9F4;border-right:4px solid #C9A227;border-radius:8px;padding:16px 20px;margin-bottom:22px;text-align:right;">'
                       f'<div style="color:#13294B;font-weight:bold;font-size:13px;">✍️ ورشة الصياغة '
                       f'<span style="color:#5B6472;font-weight:normal;font-style:italic;font-size:11px;">({_esc(lesson.get("category",""))})</span></div>'
                       f'<p style="color:#1C2430;font-size:12.5px;line-height:1.6;margin:8px 0;">{_esc(lesson.get("principle_ar",""))}</p>'
                       f'<p style="color:#8C1D1D;font-size:12px;margin:6px 0;">❌ {_esc(lesson.get("weak_example_ar",""))}</p>'
                       f'<p style="color:#1E6B3E;font-size:12px;margin:6px 0;">✅ {_esc(lesson.get("strong_example_ar",""))}</p>'
                       f'<p style="color:#5B6472;font-size:11.5px;font-style:italic;margin:6px 0 0 0;">{_esc(lesson.get("why_ar",""))}</p></div>')

    fb = content.get("flashback", {})
    flashback_html = ""
    if fb.get("content_ar"):
        pts = "".join(f'<li style="margin-bottom:4px;">{_esc(k)}</li>' for k in fb.get("key_points", []))
        body = "".join(f'<p style="color:#1C2430;font-size:14px;line-height:1.7;margin:0 0 8px 0;">{_esc(para)}</p>' for para in fb.get("content_ar", "").split("\n\n"))
        flashback_html = (f'<div style="background:#fff;border:1px solid #D8D2C2;border-radius:10px;padding:18px 20px;margin-bottom:22px;">'
                          f'<div style="color:#C9A227;font-weight:bold;font-size:12px;letter-spacing:0.5px;margin-bottom:6px;text-transform:uppercase;">🧠 ومضة إدارية — {_esc(fb.get("topic",""))}</div>'
                          f'{body}<ul style="margin:0;padding-right:18px;color:#5B6472;font-size:13px;">{pts}</ul></div>')

    term = content.get("term_of_week", {})
    term_html = ""
    if term and term.get("term_en"):
        example_html = ""
        if term.get("example_ar"):
            example_html = f'<p style="color:#E8D9A0;font-size:11px;font-style:italic;margin:8px 0 0 0;">في الممارسة: {_esc(term["example_ar"])}</p>'
        term_html = (f'<div style="background:#13294B;border-radius:10px;padding:16px 20px;margin-bottom:22px;text-align:right;">'
                     f'<div style="color:#fff;font-weight:bold;font-size:14px;">📖 {_esc(term.get("term_ar",""))} '
                     f'<span style="color:#E8D9A0;font-weight:normal;font-style:italic;font-size:11px;">({_esc(term.get("term_en",""))})</span></div>'
                     f'<p style="color:#fff;font-size:12.5px;line-height:1.6;margin:8px 0 0 0;">{_esc(term.get("definition_ar",""))}</p>{example_html}</div>')

    cards_html = ""
    for block in content.get("domain_updates", []):
        risk = block.get("risk_level", "متوسط")
        rc = RISK_COLORS.get(risk, "8A6D1A")
        ri = RISK_ICONS.get(risk, "")
        icon = DOMAIN_ICONS.get(block.get("domain", ""), "🆕")
        source_html = ""
        if block.get("source_url"):
            source_html = f'<a href="{block["source_url"]}" style="color:#13294B;font-size:12px;text-decoration:underline;">المصدر: {_esc(block.get("source_name",""))}</a>'
        skill_html = ""
        if block.get("skill_gained"):
            skill_html = f'<div style="color:#8a7a3d;font-size:10.5px;font-style:italic;margin-top:6px;">🎓 المهارة المكتسبة: {_esc(block["skill_gained"])}</div>'
        cards_html += (f'<div style="background:#fff;border:1px solid #D8D2C2;border-radius:10px;margin-bottom:20px;overflow:hidden;">'
                       f'<table role="presentation" width="100%" cellpadding="0" cellspacing="0" style="background:#13294B;"><tr>'
                       f'<td style="padding:12px 18px;text-align:right;color:#fff;font-size:15px;font-weight:bold;font-family:Calibri,Tahoma,sans-serif;">{icon} {_esc(block.get("domain",""))}</td>'
                       f'<td style="padding:12px 18px;text-align:left;white-space:nowrap;"><span style="background:#{rc};color:#fff;padding:2px 10px;border-radius:10px;font-size:11px;font-weight:bold;">{ri} {risk}</span></td>'
                       f'</tr></table><div style="padding:16px 18px;text-align:right;font-family:Calibri,Tahoma,sans-serif;">'
                       f'<div style="font-weight:bold;color:#13294B;font-size:13.5px;margin-bottom:6px;">{_esc(block.get("title",""))}</div>'
                       f'<p style="color:#1C2430;font-size:13.5px;line-height:1.7;margin:0 0 10px 0;">{_esc(block.get("summary",""))}</p>'
                       f'<div style="background:#FBF9F4;border-right:3px solid #C9A227;padding:10px 14px;border-radius:4px;">'
                       f'<div style="margin-bottom:6px;"><strong style="color:#13294B;font-size:12.5px;">الأثر:</strong> <span style="color:#334155;font-size:12.5px;">{_esc(block.get("why_it_matters",""))}</span></div>'
                       f'<div><strong style="color:#13294B;font-size:12.5px;">ماذا تفعل بهذا الآن:</strong> <span style="color:#334155;font-size:12.5px;">{_esc(block.get("recommended_actions",""))}</span></div>'
                       f'</div>{skill_html}<div style="margin-top:10px;">{source_html}</div></div></div>')

    ideas_html = ""
    for idea in content.get("engagement_ideas", []):
        qs = "".join(f'<li style="margin-bottom:4px;">{_esc(q)}</li>' for q in idea.get("scope_questions", []))
        watch = ""
        if idea.get("watch_for_ar"):
            watch = f'<p style="color:#B5541A;font-size:11px;font-style:italic;margin:6px 0 0 0;">🚩 {_esc(idea["watch_for_ar"])}</p>'
        ideas_html += (f'<div style="background:#fff;border:1px solid #D8D2C2;border-right:4px solid #13294B;border-radius:8px;padding:14px 18px;margin-bottom:14px;text-align:right;">'
                       f'<div style="color:#13294B;font-weight:bold;font-size:13px;">{_esc(idea.get("title",""))}</div>'
                       f'<p style="color:#1C2430;font-size:12px;margin:6px 0;"><strong>الهدف:</strong> {_esc(idea.get("objective_ar",""))}</p>'
                       f'<ul style="margin:4px 0;padding-right:18px;color:#334155;font-size:11.5px;">{qs}</ul>{watch}</div>')

    growth = content.get("growth_corner", {})
    wellbeing = content.get("wellbeing_corner", {})
    gw_html = ""
    if growth.get("insight_ar"):
        gw_html += (f'<div style="background:#fff;border:1px solid #D8D2C2;border-right:4px solid #13294B;border-radius:8px;padding:16px 20px;margin-bottom:14px;text-align:right;">'
                    f'<div style="color:#13294B;font-weight:bold;font-size:13px;margin-bottom:6px;">🌱 مسار الاحتراف</div>'
                    f'<p style="color:#1C2430;font-size:13px;line-height:1.6;margin:0 0 8px 0;">{_esc(growth.get("insight_ar",""))}</p>'
                    f'<p style="color:#5B6472;font-size:12px;font-style:italic;margin:0;"><strong>طبّقها هذا الأسبوع:</strong> {_esc(growth.get("how_to_apply_ar",""))}</p></div>')
    if wellbeing.get("tip_ar"):
        gw_html += (f'<div style="background:#FBF9F4;border-radius:8px;padding:14px 18px;margin-bottom:22px;text-align:right;">'
                    f'<div style="color:#13294B;font-weight:bold;font-size:13px;margin-bottom:6px;">⚖️ لحظة توازن</div>'
                    f'<p style="color:#1C2430;font-size:12.5px;line-height:1.6;margin:0;font-style:italic;">{_esc(wellbeing.get("tip_ar",""))}</p></div>')

    reading_html = "".join(
        f'<li style="margin-bottom:8px;color:#1C2430;font-size:13px;"><span style="color:#5B6472;">{_esc(it.get("type","مقال"))} — </span>'
        f'<a href="{it.get("url","#")}" style="color:#13294B;font-weight:bold;text-decoration:none;">{_esc(it.get("title",""))}</a>'
        f'<span style="color:#5B6472;font-size:12px;"> ({_esc(it.get("source_name",""))})</span></li>'
        for it in content.get("further_reading", []))
    book = content.get("book_of_week", {})
    book_html = ""
    if book.get("title"):
        takeaway = ""
        if book.get("key_takeaway_ar"):
            takeaway = f'<p style="color:#13294B;font-size:11.5px;margin:6px 0 0 0;">💡 <strong>طبّقه هذا الأسبوع:</strong> {_esc(book["key_takeaway_ar"])}</p>'
        book_html = (f'<div style="border-top:1px solid #D8D2C2;margin-top:14px;padding-top:14px;">'
                     f'<div style="color:#13294B;font-weight:bold;font-size:13.5px;">📕 كتاب الأسبوع: {_esc(book["title"])} — {_esc(book.get("author",""))}</div>'
                     f'<p style="color:#5B6472;font-size:12.5px;font-style:italic;margin:6px 0 0 0;">{_esc(book.get("why_ar",""))}</p>{takeaway}</div>')

    standards_section = ""
    if standards_html:
        standards_section = (f'<div style="text-align:right;"><div style="color:#13294B;font-weight:bold;font-size:15px;margin-bottom:14px;'
                             f'border-bottom:2px solid #C9A227;padding-bottom:6px;">🆕 رادار المعايير والإصدارات الجديدة</div>{standards_html}</div>')
    ideas_section = ""
    if ideas_html:
        ideas_section = (f'<div style="text-align:right;"><div style="color:#13294B;font-weight:bold;font-size:15px;margin-bottom:14px;'
                         f'border-bottom:2px solid #C9A227;padding-bottom:6px;">🧭 مقترحات لمهام تدقيق</div>{ideas_html}</div>')

    return f"""<!DOCTYPE html>
<html dir="rtl" lang="ar"><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0"></head>
<body style="background:#F1EFE8;margin:0;padding:20px;font-family:Calibri,Tahoma,sans-serif;">
<div style="max-width:680px;margin:0 auto;">
<div style="background:linear-gradient(135deg,#0B1F3A 0%,#13294B 100%);padding:30px 20px;border-radius:10px 10px 0 0;text-align:center;">
<div style="color:#E8D9A0;font-size:12px;letter-spacing:1px;margin-bottom:6px;">العدد {issue_no} • {datetime.now().strftime('%Y-%m-%d')}</div>
<h1 style="color:#fff;margin:0;font-size:21px;">النشرة الأسبوعية للحوكمة والتدقيق الداخلي</h1>
<div style="color:#C9A227;margin-top:8px;font-size:14px;font-style:italic;">{_esc(content.get('issue_theme',''))}</div></div>
<div style="background:#fff;padding:18px 22px;text-align:right;"><p style="color:#5B6472;font-size:13.5px;line-height:1.7;margin:0;font-style:italic;">✉️ {_esc(content.get('editor_note',''))}</p></div>
<div style="padding:20px 0 0 0;">
<div style="text-align:right;">{exec_brief_html}{stat_html}</div>
{standards_section}
<div style="text-align:right;">{lesson_html}{flashback_html}{term_html}</div>
<div style="text-align:right;"><div style="color:#13294B;font-weight:bold;font-size:15px;margin-bottom:14px;border-bottom:2px solid #C9A227;padding-bottom:6px;">📌 أبرز التطورات حسب المجال</div>{cards_html}</div>
{ideas_section}
<div style="text-align:right;">{gw_html}</div>
<div style="background:#fff;border:1px solid #D8D2C2;border-radius:10px;padding:18px 20px;text-align:right;"><div style="color:#13294B;font-weight:bold;font-size:14px;margin-bottom:10px;">📚 للقراءة هذا الأسبوع</div><ul style="margin:0;padding-right:18px;">{reading_html}</ul>{book_html}</div>
</div>
<div style="text-align:center;padding:18px;font-size:11px;color:#5B6472;border-top:1px solid #D8D2C2;margin-top:10px;">نشرة داخلية سرية — مُعدّة عبر نظام رصد الرؤى الحوكمية والتدقيقية.<br>جميع الحقوق محفوظة © {datetime.now().year}</div>
</div></body></html>"""


# ============================================================================
# إرسال الإيميل (بلا تغيير)
# ============================================================================
def send_email_report(html_content: str, report_filepath: str, issue_theme: str, issue_no: int):
    smtp_server = os.getenv("SMTP_SERVER", "smtp.gmail.com")
    smtp_port = int(os.getenv("SMTP_PORT", 465))
    sender_email = os.getenv("SENDER_EMAIL")
    sender_password = os.getenv("SENDER_PASSWORD")
    recipient_emails = [e.strip() for e in os.getenv("RECIPIENT_EMAILS", "").split(",") if e.strip()]
    if not sender_email or not sender_password or not recipient_emails:
        raise ValueError("Missing required SMTP environment variables.")

    msg = MIMEMultipart("mixed")
    msg["Subject"] = Header(f"العدد {issue_no} | {issue_theme} — {datetime.now().strftime('%Y-%m-%d')}", "utf-8")
    msg["From"], msg["To"] = sender_email
    msg.attach(MIMEText(html_content, "html", "utf-8"))

    if os.path.exists(report_filepath):
        subtype = "pdf" if report_filepath.endswith(".pdf") else "vnd.openxmlformats-officedocument.wordprocessingml.document"
        with open(report_filepath, "rb") as f:
            attachment = MIMEApplication(f.read(), _subtype=subtype)
            attachment.add_header("Content-Disposition", "attachment", filename=os.path.basename(report_filepath))
            msg.attach(attachment)
    else:
        log.warning(f"Report file missing: {report_filepath}")

    try:
        if smtp_port == 465:
            with smtplib.SMTP_SSL(smtp_server, smtp_port) as server:
                server.login(sender_email, sender_password)
                server.sendmail(sender_email, recipient_emails, msg.as_string())
        else:
            with smtplib.SMTP(smtp_server, smtp_port) as server:
                server.starttls(); server.login(sender_email, sender_password)
                server.sendmail(sender_email, recipient_emails, msg.as_string())
        log.info(f"Newsletter dispatched to {len(recipient_emails)} recipient(s).")
    except Exception as e:
        log.error(f"Failed to dispatch email: {e}", exc_info=True)
        raise


# ============================================================================
# الـ Pipeline
# ============================================================================
def run_pipeline():
    verify_fonts_installed()
    content, ok = generate_newsletter_content()
    if not ok:
        log.error("Pipeline STOPPED — no email was sent.")
        sys.exit(1)

    docx_filepath = build_word_document(content, reports_dir="reports")
    issue_no = int(re.search(r"Issue(\d+)", docx_filepath).group(1))
    report_filepath = convert_docx_to_pdf(docx_filepath)
    email_html = generate_email_html(content, issue_no)
    with open("email_preview.html", "w", encoding="utf-8") as f:
        f.write(email_html)

    send_email_report(email_html, report_filepath, content.get("issue_theme", "النشرة الأسبوعية"), issue_no)
    log.info("Pipeline complete.")

if __name__ == "__main__":
    run_pipeline()
